"""Shared helpers for filling bracketed placeholders (e.g. ``[Field Name]``)
in a .docx template while preserving the original Word formatting. Reused by
every per-document filler module (docx_filler.py, docx_filler_wad.py).
"""

from __future__ import annotations

import re
from datetime import date, datetime

from docx import Document
from docx.table import Table

PLACEHOLDER_RE = re.compile(r"\[[^\[\]]*\]")


def format_date(value: str) -> str:
    """Accepts an ISO date (YYYY-MM-DD) or a plain string and returns
    a human readable date like "July 23, 2026". Falls back to the raw
    string if it isn't a parseable ISO date."""
    if not value:
        return value
    if isinstance(value, (date, datetime)):
        return value.strftime("%B %d, %Y")
    try:
        parsed = datetime.strptime(value.strip(), "%Y-%m-%d")
        return parsed.strftime("%B %d, %Y")
    except ValueError:
        return value


def format_date_dmy(value: str) -> str:
    """Same as format_date but renders as DD-MM-YYYY."""
    if not value:
        return value
    if isinstance(value, (date, datetime)):
        return value.strftime("%d-%m-%Y")
    try:
        parsed = datetime.strptime(value.strip(), "%Y-%m-%d")
        return parsed.strftime("%d-%m-%Y")
    except ValueError:
        return value


def format_money(value: str) -> str:
    """Formats a numeric string as n'nnn'nnn.nn (apostrophe thousands
    separators, 2 decimal places). Non-numeric input is returned unchanged."""
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return text
    cleaned = text.replace(",", "").replace("'", "")
    try:
        amount = float(cleaned)
    except ValueError:
        return text
    return f"{amount:,.2f}".replace(",", "'")


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for container in containers:
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def replace_in_paragraph(paragraph, replacements: dict) -> None:
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(run.text for run in runs)
    if "[" not in full_text:
        return

    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    run_spans = []
    pos = 0
    for run in runs:
        length = len(run.text)
        run_spans.append((pos, pos + length, run))
        pos += length

    # Only known placeholders get replaced; unrecognized bracketed text
    # (e.g. field-code artifacts) is left in place.
    edits = []
    for match in matches:
        key = match.group(0)[1:-1]
        if key in replacements:
            edits.append((match.start(), match.end(), replacements[key]))

    if not edits:
        return

    # Apply right-to-left so earlier run offsets stay valid.
    for start, end, replacement in reversed(edits):
        overlapping = [
            (r_start, r_end, run)
            for r_start, r_end, run in run_spans
            if r_start < end and r_end > start
        ]
        if not overlapping:
            continue

        first_start, first_end, first_run = overlapping[0]
        last_start, last_end, last_run = overlapping[-1]

        prefix = first_run.text[: max(0, start - first_start)]
        suffix = last_run.text[max(0, end - last_start):]

        if first_run is last_run:
            first_run.text = prefix + replacement + suffix
        else:
            first_run.text = prefix + replacement
            last_run.text = suffix
            for _, _, middle_run in overlapping[1:-1]:
                middle_run.text = ""


def fill_placeholders(document: Document, replacements: dict) -> None:
    """Runs the placeholder substitution pass over every paragraph in the
    document (body, tables, headers/footers)."""
    for paragraph in iter_all_paragraphs(document):
        replace_in_paragraph(paragraph, replacements)
