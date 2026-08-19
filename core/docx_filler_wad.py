"""Fills bracketed placeholders in the WAD (Working Advance) template.

The settlement table has exactly one pre-formatted entry row followed by
12 blank rows before the "Total" row, so up to 13 SWA (Settlement of
Working Advance) entries are supported without altering the table's
structure. The first entry is filled via the normal placeholder
substitution pass (its brackets are the only occurrence of ``[SWA Date]``
etc. in the document); any further entries are written directly into the
template's blank rows.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from .docx_utils import fill_placeholders, format_date, format_date_dmy, format_money

STAFF_PER_DIEM_PREFIX = "Staff Per Diem"
MAX_SWA_ENTRIES = 13

# Scalar fields the caller must supply.
SCALAR_REQUIRED_FIELDS = [
    "WAD Number",
    "WAD Amount",
    "WAD Amount in Words",
    "WAD Purpose",
    "WAD Date",
    "HWA Date",
]

# Scalar fields that may be left blank.
SCALAR_OPTIONAL_FIELDS = ["SWA Received", "SWA Returned"]

# Per-entry fields; SWA Date/Receipt Number are conditionally required.
ENTRY_REQUIRED_FIELDS = ["SWA Detail", "SWA Expenses"]
ENTRY_CONDITIONAL_FIELDS = ["SWA Date", "SWA Receipt Number"]

_SETTLEMENT_TABLE_INDEX = 1
_FIRST_BLANK_ROW_INDEX = 4  # row 3 holds entry #1's placeholders


def _parse_amount(value, field_name: str) -> float:
    text = str(value or "").strip().replace(",", "").replace("'", "")
    if not text:
        return 0.0
    try:
        return float(text)
    except ValueError:
        raise ValueError(f"{field_name} must be a number")


def _is_staff_per_diem(detail: str) -> bool:
    return detail.strip().startswith(STAFF_PER_DIEM_PREFIX)


def _validate_entries(entries: list) -> None:
    if not isinstance(entries, list) or len(entries) < 1:
        raise ValueError("At least one SWA entry is required")
    if len(entries) > MAX_SWA_ENTRIES:
        raise ValueError(f"A maximum of {MAX_SWA_ENTRIES} SWA entries is supported")

    for index, entry in enumerate(entries, start=1):
        detail = str(entry.get("SWA Detail", "")).strip()
        expenses = str(entry.get("SWA Expenses", "")).strip()
        missing = []
        if not detail:
            missing.append("SWA Detail")
        if not expenses:
            missing.append("SWA Expenses")
        if not _is_staff_per_diem(detail):
            if not str(entry.get("SWA Date", "")).strip():
                missing.append("SWA Date")
            if not str(entry.get("SWA Receipt Number", "")).strip():
                missing.append("SWA Receipt Number")
        if missing:
            raise ValueError(f"SWA entry {index}: missing {', '.join(missing)}")


def prepare_wad_data(data: dict):
    """Validates and normalizes raw request data. Returns (values, extra_rows)
    where ``values`` is the placeholder -> text map for the substitution pass
    (covering all scalar fields plus SWA entry #1), and ``extra_rows`` is the
    list of already-formatted {date, receipt, detail, expenses} dicts for any
    SWA entries beyond the first, to be written directly into the template's
    blank rows."""
    missing = [f for f in SCALAR_REQUIRED_FIELDS if not str(data.get(f, "")).strip()]
    if missing:
        raise ValueError(f"Missing required field(s): {', '.join(missing)}")

    entries = data.get("SWA Entries") or []
    _validate_entries(entries)

    wad_amount = _parse_amount(data.get("WAD Amount"), "WAD Amount")
    swa_received_raw = str(data.get("SWA Received", "")).strip()
    swa_received_amount = _parse_amount(swa_received_raw, "SWA Received") if swa_received_raw else 0.0
    swa_returned_raw = str(data.get("SWA Returned", "")).strip()
    swa_returned_amount = _parse_amount(swa_returned_raw, "SWA Returned") if swa_returned_raw else 0.0
    swa_amount_total = sum(
        _parse_amount(entry.get("SWA Expenses"), "SWA Expenses") for entry in entries
    )
    total_income = wad_amount + swa_received_amount
    total_expense = swa_amount_total + swa_returned_amount

    values = {
        "WAD Number": str(data["WAD Number"]).strip(),
        "WAD Amount": format_money(data["WAD Amount"]),
        "WAD Amount in Words": str(data["WAD Amount in Words"]).strip(),
        "WAD Purpose": str(data["WAD Purpose"]).strip(),
        "WAD Date": format_date(str(data["WAD Date"]).strip()),
        "HWA Date": format_date_dmy(str(data["HWA Date"]).strip()),
        "SWA Amount Total": format_money(str(swa_amount_total)),
        "Total Income": format_money(str(total_income)),
        "Total Expense": format_money(str(total_expense)),
        "SWA Received": f"PHP {format_money(swa_received_raw)}" if swa_received_raw else "",
        "SWA Returned": format_money(swa_returned_raw) if swa_returned_raw else "0.00",
    }

    def _formatted_row(entry: dict) -> dict:
        date_raw = str(entry.get("SWA Date", "")).strip()
        return {
            "date": format_date_dmy(date_raw) if date_raw else "",
            "receipt": str(entry.get("SWA Receipt Number", "")).strip(),
            "detail": str(entry.get("SWA Detail", "")).strip(),
            "expenses": format_money(str(entry.get("SWA Expenses", "")).strip()),
        }

    first_row = _formatted_row(entries[0])
    values["SWA Date"] = first_row["date"]
    values["SWA Receipt Number"] = first_row["receipt"]
    values["SWA Detail"] = first_row["detail"]
    values["SWA Expenses"] = first_row["expenses"]

    extra_rows = [_formatted_row(entry) for entry in entries[1:]]

    return values, extra_rows


def _set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run.text = ""
    if text:
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(8)


def _fill_extra_swa_rows(document: Document, extra_rows: list) -> None:
    table = document.tables[_SETTLEMENT_TABLE_INDEX]
    for offset, row_values in enumerate(extra_rows):
        cells = table.rows[_FIRST_BLANK_ROW_INDEX + offset].cells
        _set_cell_text(cells[0], row_values["date"])
        _set_cell_text(cells[1], row_values["receipt"])
        _set_cell_text(cells[2], row_values["detail"])
        _set_cell_text(cells[3], row_values["detail"])
        _set_cell_text(cells[5], row_values["expenses"])


def fill_template_wad(template_path: str, output_path: str, data: dict) -> str:
    """Fills the WAD template with the given data and writes the result to
    output_path. Returns output_path."""
    values, extra_rows = prepare_wad_data(data)

    document = Document(template_path)
    fill_placeholders(document, values)
    if extra_rows:
        _fill_extra_swa_rows(document, extra_rows)

    document.save(output_path)
    return output_path
